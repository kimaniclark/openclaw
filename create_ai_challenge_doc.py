from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('Dan Martell AI CEO Challenge', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('4-Day Challenge - March 9-12, 2026')
doc.add_paragraph('From: dm@danmartell.com')
doc.add_paragraph('')

# Day 1
doc.add_heading('Day 1: ChatGPT isn\'t cutting it (what to do instead)', level=1)
doc.add_paragraph('Date: Monday, March 9, 2026')
doc.add_paragraph('')

day1_content = """Hey Kimani,

Challenge Day 1 is live. Let's go 👊

Let me guess how you've been using AI:

You open ChatGPT. You ask it to write something… It sounds generic. You ask it to research something… It's not quite accurate. Felt like magic at first, but now you're seeing the holes.

That's because everything has changed. Again.

Becoming an AI CEO isn't about one tool. It's a toolbox.

Most CEOs are using a hammer for every job. That ends today.

Here's how AI CEOs actually operate in 2026:

→ They open ChatGPT voice mode mid-meeting
→ After the meeting, they send Grok to go analyze trends
→ Then Gemini starts scanning YouTube for inspiration
→ Then they take all of that into Claude to finish the final product

Right tool. Right job. Every time.

That's what you and 30,000 others are doing Day 1 of this challenge. Building your AI stack.

BUT… the problem then becomes prompting each one properly.

Here's the good news:

I'm walking you through the one method that works across ALL these tools. It's called Reverse Prompting.

You'll find:
• My 5-tool core stack (with real examples)
• The Reverse Prompt formula you can copy right now
• My complete 2026 AI stack: 36 tools I actually use in my businesses

By the end of today, you'll stop using AI like it's 2023.

That's the foundation for everything we're building this week.

-DM

P.S. Message me "BONUS" on Instagram to get my full AI Tool Stack doc. @danmartell 👊
"""
doc.add_paragraph(day1_content)

# Day 2
doc.add_heading('Day 2: Your Digital Brain', level=1)
doc.add_paragraph('Date: Tuesday, March 10, 2026')
doc.add_paragraph('')

day2_content = """Hey Kimani,

Day 2 just dropped.

Yesterday you learned the AI stack. Today you're solving the problem everyone hits next.

The re-explaining problem.

Every new AI tool = starting from scratch. Explaining yourself. Your business. Your priorities. Over and over.

Day 2 fixes that.

You're building your Digital Brain. One document. Paste it anywhere. Instantly known.

Takes 15 minutes. Portable forever.

See you inside.

-DM

P.S. If you missed the bonus yesterday (so many people reached out I think I broke the messaging limits 🤯 some might not have gone through)… just hit me up again: DM me "Bonus" on Instagram.
"""
doc.add_paragraph(day2_content)

# Day 3
doc.add_heading('Day 3: AI Agents (This Changes Everything)', level=1)
doc.add_paragraph('Date: Wednesday, March 11, 2026')
doc.add_paragraph('')

day3_content = """Hey Kimani,

AI Challenge Day 3 is live!

Days 1 and 2 were setup. Today is the payoff.

You learned the AI stack. You built your Digital Brain.

Today you're handing off entire tasks and walking away.

Not "help me think about this."

Do it. Come back when it's done.

This is the shift most people never make.

See you inside.

-DM

P.S. Tomorrow is Day 4: Your 12-month roadmap + a major announcement I've been holding back.
"""
doc.add_paragraph(day3_content)

# Day 4
doc.add_heading('Day 4: The Roadmap (and the thing I\'ve been holding back)', level=1)
doc.add_paragraph('Date: Thursday, March 12, 2026')
doc.add_paragraph('')

day4_content = """Hey Kimani,

This is it.

Day 4 is where it all comes together.

I showed you how to build one AI agent from scratch that saves you hours every week.

Today I'm showing you the roadmap for the other 20+ areas.

What to build next. In what order. How to scale to 20+ hours saved per week.

But here's the reality…

You can take what you learned and run with it yourself. Follow the roadmap. Build more agents. And it'll help.

But here's what I've seen over and over:

AI systems start falling apart if they're not built around SOLID business systems.

AI is a multiplier. It will amplify whatever you're already doing.

So if your business systems (time systems, team systems, sales, marketing) are weak, messy, or not followed… AI will just create more chaos.

All those tools and applications become headaches instead of fuel.

Which is why you have to build both at the same time.

Business systems and AI systems.

This is exactly what I teach inside my Elite Coaching program.

Here's what you get:
• Weekly calls where I coach you and help you implement, set vision, and find roadblocks.
• The complete system - Elite Growth Model, Achievement Roadmap, Rocket Selling System, AI Adoption, plus 50+ more.
• Community of 1,500+ entrepreneurs building alongside you who've already made this shift.

Either way - seriously, congrats on finishing.

You're ahead of 90% of CEOs right now.

Now decide: How far ahead do you want to be in 12 months?

-DM
"""
doc.add_paragraph(day4_content)

# Notes section
doc.add_heading('Notes & Resources', level=1)
doc.add_paragraph('')
notes = """Key Takeaways:

DAY 1 - AI Stack & Reverse Prompting:
• Build a toolbox, not rely on one tool
• 5-tool core stack: ChatGPT, Grok, Gemini, Claude, + 1 more
• Reverse Prompting formula for better results
• Full AI stack: 36 tools

DAY 2 - Digital Brain:
• One document that explains you/your business
• Paste into any AI tool for instant context
• 15 minutes to create, portable forever

DAY 3 - AI Agents:
• Hand off entire tasks, not just questions
• "Do it. Come back when it's done."
• The shift from AI assistant → AI worker

DAY 4 - Roadmap:
• 20+ areas to build AI agents
• Order of implementation matters
• AI amplifies existing business systems (good or bad)
• Build business systems AND AI systems together

BONUS: DM @danmartell "BONUS" on Instagram for AI Tool Stack doc
"""
doc.add_paragraph(notes)

# Save
doc.save('Dan-Martell-AI-Challenge-Notes.docx')
print("✅ Word document created: Dan-Martell-AI-Challenge-Notes.docx")

